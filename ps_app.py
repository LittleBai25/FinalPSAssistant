import streamlit as st
import os
import io
from PIL import Image
from datetime import datetime
import uuid
import base64
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import asyncio

# Import custom modules
from agents.ps_info_collector import PSInfoCollector
from agents.supporting_file_analyzer import SupportingFileAnalyzer
from agents.ps_analyzer import PSAnalyzer
from agents.ps_rewriter import PSRewriter
from config.prompts import load_prompts, save_prompts

# Set page configuration
st.set_page_config(
    page_title="PS Assistant Tool",
    page_icon="📝",
    layout="wide"
)

# Initialize session state
if "university_info_report" not in st.session_state:
    st.session_state.university_info_report = None
if "supporting_file_analysis" not in st.session_state:
    st.session_state.supporting_file_analysis = None
if "ps_analysis_strategy" not in st.session_state:
    st.session_state.ps_analysis_strategy = None
if "rewritten_ps" not in st.session_state:
    st.session_state.rewritten_ps = None
if "ps_content" not in st.session_state:
    st.session_state.ps_content = None
if "university" not in st.session_state:
    st.session_state.university = ""
if "major" not in st.session_state:
    st.session_state.major = ""
if "custom_requirements" not in st.session_state:
    st.session_state.custom_requirements = ""
if "supporting_files" not in st.session_state:
    st.session_state.supporting_files = None
if "ps_file" not in st.session_state:
    st.session_state.ps_file = None
if "current_step" not in st.session_state:
    st.session_state.current_step = 1
if "info_collector_model" not in st.session_state:
    st.session_state.info_collector_model = "anthropic/claude-3-7-sonnet"
if "supporting_analyzer_model" not in st.session_state:
    st.session_state.supporting_analyzer_model = "anthropic/claude-3-7-sonnet"
if "ps_analyzer_model" not in st.session_state:
    st.session_state.ps_analyzer_model = "anthropic/claude-3-7-sonnet"
if "ps_rewriter_model" not in st.session_state:
    st.session_state.ps_rewriter_model = "anthropic/claude-3-7-sonnet"

# 支持的模型列表
SUPPORTED_MODELS = [
    "anthropic/claude-3-7-sonnet",
    "anthropic/claude-3-5-sonnet",
    "openai/gpt-4-turbo",
    "openai/gpt-4-1106-preview",
    "qwen/qwen-max"
]

# 检查必要的API密钥是否设置
def check_api_keys():
    """检查Streamlit secrets中是否设置了必要的API密钥。"""
    api_keys = {
        "OPENROUTER_API_KEY": st.secrets.get("OPENROUTER_API_KEY", None),
        "SERPER_API_KEY": st.secrets.get("SERPER_API_KEY", None)
    }
    
    return {k: bool(v) for k, v in api_keys.items()}

# 创建Word文档报告并提供下载
def create_downloadable_report(report_title, report_content):
    """生成可下载的Word文档报告"""
    # 创建一个新的Word文档
    doc = docx.Document()
    
    # 设置文档标题
    title = doc.add_heading(report_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d')}")
    date_run.italic = True
    
    # 添加分隔线
    doc.add_paragraph("_" * 50)
    
    # 添加报告内容 (处理Markdown)
    # 这是一个简单的处理，实际应用中可能需要更复杂的Markdown转换
    lines = report_content.split('\n')
    current_paragraph = None
    
    for line in lines:
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # 处理列表项
        elif line.startswith('- ') or line.startswith('* '):
            if not current_paragraph or not current_paragraph.text.startswith(('- ', '* ')):
                current_paragraph = doc.add_paragraph()
            current_paragraph.add_run('\n' + line)
        # 处理空行
        elif not line.strip():
            current_paragraph = None
        # 处理普通文本
        else:
            if not current_paragraph:
                current_paragraph = doc.add_paragraph()
            current_paragraph.add_run(line)
    
    # 保存文档到内存中
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    # 转换为Base64编码以便于下载
    base64_docx = base64.b64encode(docx_stream.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_docx}" download="{report_title}.docx">点击下载Word文档</a>'
    
    return href

# Main function
def main():
    # Create tabs
    tab1, tab2, tab3 = st.tabs(["PS处理助手", "提示词调试", "系统状态"])
    
    with tab1:
        st.title("PS分稿测试助手")
        
        # 步骤1：输入基本信息并获取院校信息报告
        if st.session_state.current_step == 1:
            st.subheader("步骤1：输入院校专业信息")
            
            # 输入院校和专业信息
            university = st.text_input("目标院校", value=st.session_state.university)
            st.session_state.university = university
            
            major = st.text_input("目标专业", value=st.session_state.major)
            st.session_state.major = major
            
            # 生成院校信息报告按钮 (只有必填字段已填写时启用)
            generate_enabled = university and major
            
            if not generate_enabled:
                st.info("请输入目标院校和专业以生成信息收集报告。")
            
            col1, col2 = st.columns([3, 1])
            with col2:
                if st.button("生成院校信息报告", disabled=not generate_enabled, key="generate_uni_info", use_container_width=True):
                    if university and major:
                        with st.spinner(f"使用 {st.session_state.info_collector_model} 收集院校专业信息..."):
                            # 创建院校信息收集代理
                            info_collector = PSInfoCollector(model_name=st.session_state.info_collector_model)
                            
                            # 收集院校信息
                            university_info = asyncio.run(info_collector.collect_information(
                                university=university,
                                major=major,
                                custom_requirements="" # 保留参数但传入空字符串
                            ))
                            
                            # 保存院校信息报告
                            st.session_state.university_info_report = university_info
                            
                            # 更新步骤
                            st.session_state.current_step = 2
                            st.rerun()
        
        # 步骤2：上传支持文件和PS初稿，生成分析报告
        elif st.session_state.current_step == 2:
            # 显示院校信息报告
            with st.expander("院校信息收集报告", expanded=True):
                st.markdown(st.session_state.university_info_report)
                
                # 添加导出Word文档的按钮
                report_download = create_downloadable_report(
                    f"{st.session_state.university} {st.session_state.major}专业信息收集报告",
                    st.session_state.university_info_report
                )
                st.markdown(report_download, unsafe_allow_html=True)
            
            # 添加返回按钮
            if st.button("返回院校搜索", key="return_to_search"):
                st.session_state.current_step = 1
                st.rerun()
            
            st.markdown("---")
            
            st.subheader("步骤2：上传支持文件和PS初稿")
            
            # 添加写作需求输入
            writing_requirements = st.text_area(
                "写作需求",
                value=st.session_state.custom_requirements,
                placeholder="输入你的PS写作需求或特殊要求..."
            )
            st.session_state.custom_requirements = writing_requirements
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 支持文件上传")
                st.markdown("上传简历、成绩单等支持文件，帮助分析PS改写策略。")
                
                # 支持文件上传（多文件）
                supporting_files = st.file_uploader(
                    "上传支持文件（可选，可多选）",
                    type=["pdf", "jpg", "jpeg", "png", "txt"],
                    accept_multiple_files=True
                )
                
                # 处理支持文件上传
                if supporting_files:
                    st.session_state.supporting_files = supporting_files
                    st.success(f"已上传 {len(supporting_files)} 个支持文件。")
            
            with col2:
                st.markdown("### PS初稿上传")
                st.markdown("上传你的PS初稿文件，系统将基于院校信息和支持文件进行分析。")
                
                # PS初稿上传
                ps_file = st.file_uploader(
                    "上传PS初稿（必需）",
                    type=["pdf", "doc", "docx", "txt"]
                )
                
                # 处理PS初稿上传
                if ps_file:
                    st.session_state.ps_file = ps_file
                    st.success(f"已上传PS初稿：{ps_file.name}")
            
            # 生成分析报告按钮
            analyze_enabled = st.session_state.ps_file is not None
            
            if not analyze_enabled:
                st.info("请上传PS初稿文件以生成改写策略。")
            
            col1, col2 = st.columns([3, 1])
            with col2:
                if st.button("生成PS改写策略", disabled=not analyze_enabled, key="generate_ps_strategy", use_container_width=True):
                    if st.session_state.ps_file:
                        # 首先分析支持文件（如果有）
                        if st.session_state.supporting_files:
                            with st.spinner(f"使用 {st.session_state.supporting_analyzer_model} 分析支持文件..."):
                                # 创建支持文件分析代理
                                supporting_analyzer = SupportingFileAnalyzer(model_name=st.session_state.supporting_analyzer_model)
                                
                                # 分析支持文件
                                supporting_analysis = supporting_analyzer.analyze_files(
                                    uploaded_files=st.session_state.supporting_files
                                )
                                
                                # 保存支持文件分析报告
                                st.session_state.supporting_file_analysis = supporting_analysis
                        else:
                            st.session_state.supporting_file_analysis = "未提供支持文件，跳过支持文件分析环节。"
                        
                        # 然后分析PS初稿
                        with st.spinner(f"使用 {st.session_state.ps_analyzer_model} 分析PS初稿..."):
                            # 创建PS分析代理
                            ps_analyzer = PSAnalyzer(model_name=st.session_state.ps_analyzer_model)
                            
                            # 提取PS文件内容
                            ps_file = st.session_state.ps_file
                            ps_analyzer = PSAnalyzer(model_name=st.session_state.ps_analyzer_model)
                            ps_content = ps_analyzer._extract_ps_content(ps_file)
                            st.session_state.ps_content = ps_content
                            
                            # 分析PS初稿，传递写作需求
                            ps_strategy = ps_analyzer.analyze_ps(
                                ps_file=st.session_state.ps_file,
                                university_info=st.session_state.university_info_report,
                                supporting_file_analysis=st.session_state.supporting_file_analysis,
                                writing_requirements=st.session_state.custom_requirements # 传递写作需求
                            )
                            
                            # 保存PS分析策略报告
                            st.session_state.ps_analysis_strategy = ps_strategy
                            
                            # 更新步骤
                            st.session_state.current_step = 3
                            st.rerun()
        
        # 步骤3：展示PS改写策略并执行改写
        elif st.session_state.current_step == 3:
            # 显示院校信息报告（可折叠）
            with st.expander("院校信息收集报告", expanded=False):
                st.markdown(st.session_state.university_info_report)
                
                # 添加导出Word文档的按钮
                info_report_download = create_downloadable_report(
                    f"{st.session_state.university} {st.session_state.major}专业信息收集报告",
                    st.session_state.university_info_report
                )
                st.markdown(info_report_download, unsafe_allow_html=True)
            
            # 如果有支持文件分析，则显示（可折叠）
            if st.session_state.supporting_file_analysis != "未提供支持文件，跳过支持文件分析环节。":
                with st.expander("支持文件分析报告", expanded=False):
                    st.markdown(st.session_state.supporting_file_analysis)
                    
                    # 添加导出Word文档的按钮
                    supporting_download = create_downloadable_report(
                        "支持文件分析报告",
                        st.session_state.supporting_file_analysis
                    )
                    st.markdown(supporting_download, unsafe_allow_html=True)
            
            # 显示PS改写策略报告
            with st.expander("PS改写策略报告", expanded=True):
                st.markdown(st.session_state.ps_analysis_strategy)
                
                # 添加导出Word文档的按钮
                strategy_download = create_downloadable_report(
                    "PS改写策略报告",
                    st.session_state.ps_analysis_strategy
                )
                st.markdown(strategy_download, unsafe_allow_html=True)
            
            st.markdown("---")
            
            st.subheader("步骤3：PS改写")
            
            # 执行改写按钮
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col2:
                # 重置按钮（清除所有结果）
                if st.button("重新开始", key="reset_analysis", use_container_width=True):
                    # 清空所有会话状态
                    st.session_state.university_info_report = None
                    st.session_state.supporting_file_analysis = None
                    st.session_state.ps_analysis_strategy = None
                    st.session_state.rewritten_ps = None
                    st.session_state.ps_content = None
                    st.session_state.supporting_files = None
                    st.session_state.ps_file = None
                    st.session_state.current_step = 1
                    # 重新加载页面
                    st.rerun()
            
            with col3:
                if st.session_state.rewritten_ps is None:
                    if st.button("开始改写PS", key="rewrite_ps", use_container_width=True):
                        with st.spinner(f"使用 {st.session_state.ps_rewriter_model} 改写PS..."):
                            # 创建PS改写代理
                            ps_rewriter = PSRewriter(model_name=st.session_state.ps_rewriter_model)
                            
                            # 执行PS改写
                            rewritten_ps = ps_rewriter.rewrite_ps(
                                ps_content=st.session_state.ps_content,
                                rewrite_strategy=st.session_state.ps_analysis_strategy,
                                university_info=st.session_state.university_info_report
                            )
                            
                            # 保存改写后的PS
                            st.session_state.rewritten_ps = rewritten_ps
                            
                            # 重新加载页面
                            st.rerun()
            
            # 显示改写后的PS
            if st.session_state.rewritten_ps:
                st.markdown("### 改写后的PS")
                
                # 显示改写前后对比
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### 原始PS")
                    st.text_area("", st.session_state.ps_content, height=400, disabled=True)
                
                with col2:
                    st.markdown("#### 改写后的PS")
                    st.text_area("", st.session_state.rewritten_ps, height=400)
                
                # 添加导出Word文档的按钮
                rewritten_download = create_downloadable_report(
                    "改写后的Personal Statement",
                    st.session_state.rewritten_ps
                )
                st.markdown(rewritten_download, unsafe_allow_html=True)
    
    with tab2:
        st.title("模型与提示词配置")
        
        # 添加模型选择到提示词调试页面顶部
        st.subheader("模型选择")
        col1, col2 = st.columns(2)
        
        with col1:
            # 院校信息收集模型选择
            info_collector_model = st.selectbox(
                "选择院校信息收集模型",
                SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.info_collector_model) if st.session_state.info_collector_model in SUPPORTED_MODELS else 0,
                key="info_collector_model_select"
            )
            st.session_state.info_collector_model = info_collector_model
            
            # PS分析模型选择
            ps_analyzer_model = st.selectbox(
                "选择PS分析模型",
                SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.ps_analyzer_model) if st.session_state.ps_analyzer_model in SUPPORTED_MODELS else 0,
                key="ps_analyzer_model_select"
            )
            st.session_state.ps_analyzer_model = ps_analyzer_model
            
        with col2:
            # 支持文件分析模型选择
            supporting_analyzer_model = st.selectbox(
                "选择支持文件分析模型",
                SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.supporting_analyzer_model) if st.session_state.supporting_analyzer_model in SUPPORTED_MODELS else 0,
                key="supporting_analyzer_model_select"
            )
            st.session_state.supporting_analyzer_model = supporting_analyzer_model
            
            # PS改写模型选择
            ps_rewriter_model = st.selectbox(
                "选择PS改写模型",
                SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.ps_rewriter_model) if st.session_state.ps_rewriter_model in SUPPORTED_MODELS else 0,
                key="ps_rewriter_model_select"
            )
            st.session_state.ps_rewriter_model = ps_rewriter_model
        
        # 添加模型选择说明
        st.info("这些模型设置将应用于相应的处理环节。您的选择将保存在会话中。")
        
        st.markdown("---")
        
        # 加载当前提示词
        prompts = load_prompts()
        
        st.subheader("院校信息收集代理 (Agent 1)")
        
        info_collector_role = st.text_area("角色描述", prompts["ps_info_collector"]["role"], height=150)
        info_collector_task = st.text_area("任务描述", prompts["ps_info_collector"]["task"], height=200)
        info_collector_output = st.text_area("输出格式", prompts["ps_info_collector"]["output"], height=200)
        
        st.subheader("支持文件分析代理 (Agent 2.1)")
        
        supporting_analyzer_role = st.text_area("角色描述", prompts["supporting_file_analyzer"]["role"], height=150)
        supporting_analyzer_task = st.text_area("任务描述", prompts["supporting_file_analyzer"]["task"], height=200)
        supporting_analyzer_output = st.text_area("输出格式", prompts["supporting_file_analyzer"]["output"], height=200)
        
        st.subheader("PS分析代理 (Agent 2.2)")
        
        ps_analyzer_role = st.text_area("角色描述", prompts["ps_analyzer"]["role"], height=150)
        ps_analyzer_task = st.text_area("任务描述", prompts["ps_analyzer"]["task"], height=200)
        ps_analyzer_output = st.text_area("输出格式", prompts["ps_analyzer"]["output"], height=200)
        
        st.subheader("PS改写代理 (Agent 3)")
        
        ps_rewriter_role = st.text_area("角色描述", prompts["ps_rewriter"]["role"], height=150)
        ps_rewriter_task = st.text_area("任务描述", prompts["ps_rewriter"]["task"], height=200)
        ps_rewriter_output = st.text_area("输出格式", prompts["ps_rewriter"]["output"], height=150)
        
        # 保存按钮
        if st.button("保存提示词"):
            # 更新提示词字典
            prompts["ps_info_collector"]["role"] = info_collector_role
            prompts["ps_info_collector"]["task"] = info_collector_task
            prompts["ps_info_collector"]["output"] = info_collector_output
            
            prompts["supporting_file_analyzer"]["role"] = supporting_analyzer_role
            prompts["supporting_file_analyzer"]["task"] = supporting_analyzer_task
            prompts["supporting_file_analyzer"]["output"] = supporting_analyzer_output
            
            prompts["ps_analyzer"]["role"] = ps_analyzer_role
            prompts["ps_analyzer"]["task"] = ps_analyzer_task
            prompts["ps_analyzer"]["output"] = ps_analyzer_output
            
            prompts["ps_rewriter"]["role"] = ps_rewriter_role
            prompts["ps_rewriter"]["task"] = ps_rewriter_task
            prompts["ps_rewriter"]["output"] = ps_rewriter_output
            
            # 保存更新后的提示词
            save_prompts(prompts)
            st.success("提示词已成功保存！")

    with tab3:
        st.title("系统状态")
        
        # 检查API密钥
        api_key_status = check_api_keys()
        
        st.subheader("API密钥")
        
        # 显示API密钥状态表格
        status_data = [
            {"API密钥": key, "状态": "✅ 已设置" if status else "❌ 未设置"} 
            for key, status in api_key_status.items()
        ]
        
        st.table(status_data)
        
        # 添加依赖项信息
        st.subheader("依赖项信息")
        
        st.markdown("""
        **必要依赖项：**
        - `streamlit`: UI界面
        - `PyMuPDF` (fitz): PDF文件处理
        - `python-docx`: Word文档处理
        - `markitdown`: DOC格式处理
        - `Pillow` (PIL): 图像处理
        
        **安装命令：**
        ```
        pip install streamlit pymupdf python-docx pillow markitdown
        ```
        """)
        
        # 添加使用说明
        st.subheader("使用说明")
        
        st.markdown("""
        **PS分稿测试助手工作流程：**
        
        1. **步骤1 - 院校信息收集：**
           - 输入目标院校和专业信息
           - 系统搜索并整理相关申请要求和项目信息
           - 生成院校信息收集报告
        
        2. **步骤2 - 文件分析：**
           - 上传支持文件（简历、成绩单等，可选）
           - 上传PS初稿（必需）
           - 系统分析支持文件和PS初稿
           - 生成PS改写策略报告
        
        3. **步骤3 - PS改写：**
           - 查看PS改写策略报告
           - 点击"开始改写PS"按钮
           - 系统根据策略改写PS
           - 展示和下载改写后的PS
        
        所有报告均可下载为Word文档保存。
        """)

# Run the app
if __name__ == "__main__":
    main() 